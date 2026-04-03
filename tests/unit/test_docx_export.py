"""Contract tests for DOCX export (P04-03).

AC-1: DOCX export generated from selected version.
AC-2: Includes section headers, IDs, logic notes.
AC-3: Export artifact stored in project Outputs with provenance.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document

from packages.shared.assistant_context import (
    AssistantContext,
    BriefContext,
    Methodology,
    WorkflowStage,
)
from packages.shared.draft_config import DraftStore
from packages.shared.questionnaire_schema import (
    Question,
    QuestionType,
    Questionnaire,
    ResponseOption,
    Section,
)
from packages.exporters.docx_export import ExportArtifact, export_questionnaire_docx
from packages.survey_generation.engine import generate_questionnaire

NOW = datetime.now(tz=timezone.utc)


def _generate() -> Questionnaire:
    store = DraftStore()
    draft = store.create("proj-001", Methodology.SEGMENTATION)
    ctx = AssistantContext(
        project_id="proj-001",
        stage=WorkflowStage.QUESTIONNAIRE,
        methodology=Methodology.SEGMENTATION,
        brief=BriefContext(
            brief_id="brief-001", objectives="Test", audience="Adults",
            category="snack bars", geography="US", uploaded_at=NOW,
        ),
        selected_sections=["screener", "attitudes", "demographics"],
    )
    return generate_questionnaire(draft, ctx)


def _read_docx_text(content: bytes) -> str:
    """Extract all text from a DOCX byte string."""
    import io
    doc = Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# AC-1: DOCX export generated from selected version
# ---------------------------------------------------------------------------

class TestDocxGeneration:
    def test_export_returns_artifact(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        assert isinstance(artifact, ExportArtifact)

    def test_artifact_has_content(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        assert len(artifact.content) > 0

    def test_artifact_is_valid_docx(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        import io
        doc = Document(io.BytesIO(artifact.content))
        assert len(doc.paragraphs) > 0

    def test_filename_includes_project_and_version(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        assert "proj-001" in artifact.filename
        assert "v1" in artifact.filename
        assert artifact.filename.endswith(".docx")

    def test_export_different_versions(self):
        qre = _generate()
        a1 = export_questionnaire_docx(qre)
        qre.version = 2
        a2 = export_questionnaire_docx(qre)
        assert "v1" in a1.filename
        assert "v2" in a2.filename


# ---------------------------------------------------------------------------
# AC-2: Includes section headers, IDs, logic notes
# ---------------------------------------------------------------------------

class TestDocxContent:
    def test_contains_section_headers(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "Screener" in text
        assert "Demographics" in text

    def test_contains_question_ids(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "SCR_01" in text
        assert "DEM_01" in text

    def test_contains_variable_names(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "var:" in text

    def test_contains_question_text(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        # Should contain actual question text from generated questionnaire
        assert "?" in text or "snack bar" in text.lower()

    def test_contains_response_options(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        # Screener has response options like "Daily", "Weekly"
        assert "Daily" in text or "Yes" in text

    def test_contains_termination_markers(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "TERMINATE" in text

    def test_contains_scale_labels(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "Strongly Disagree" in text or "Strongly Agree" in text

    def test_contains_question_type(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "Single Select" in text or "Likert Scale" in text

    def test_contains_methodology(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "segmentation" in text.lower()

    def test_logic_notes_included_when_present(self):
        qre = Questionnaire(
            project_id="proj-001", methodology="segmentation",
            sections=[
                Section(section_id="s1", section_type="test", label="Test", order=0,
                        questions=[
                            Question(question_id="Q1", question_text="Q?",
                                     question_type=QuestionType.OPEN_ENDED, var_name="Q1",
                                     logic="Show only if SCR_01 == 1"),
                        ]),
            ],
        )
        artifact = export_questionnaire_docx(qre)
        text = _read_docx_text(artifact.content)
        assert "Logic:" in text
        assert "SCR_01" in text


# ---------------------------------------------------------------------------
# AC-3: Export artifact stored with provenance
# ---------------------------------------------------------------------------

class TestArtifactProvenance:
    def test_provenance_has_required_fields(self):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        prov = artifact.provenance()
        assert prov["questionnaire_id"] == qre.questionnaire_id
        assert prov["version"] == qre.version
        assert prov["format"] == "docx"
        assert prov["size_bytes"] > 0
        assert prov["filename"].endswith(".docx")
        assert prov["created_at"]

    def test_save_to_directory(self, tmp_path: Path):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        outputs_dir = tmp_path / "Outputs"
        saved_path = artifact.save_to(outputs_dir)
        assert saved_path.exists()
        assert saved_path.stat().st_size == artifact.size_bytes
        assert saved_path.name == artifact.filename

    def test_save_creates_directory(self, tmp_path: Path):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        nested = tmp_path / "project" / "Outputs" / "exports"
        saved_path = artifact.save_to(nested)
        assert saved_path.exists()

    def test_saved_file_is_valid_docx(self, tmp_path: Path):
        qre = _generate()
        artifact = export_questionnaire_docx(qre)
        saved_path = artifact.save_to(tmp_path)
        doc = Document(str(saved_path))
        assert len(doc.paragraphs) > 0
