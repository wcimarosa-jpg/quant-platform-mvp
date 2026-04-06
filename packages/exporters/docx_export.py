"""DOCX questionnaire export.

Generates a styled Word document from a Questionnaire model.
Includes section headers, question IDs, variable names,
response options, scale labels, and logic notes.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from packages.shared.questionnaire_schema import (
    Question,
    QuestionType,
    Questionnaire,
    ResponseOption,
    Section,
)


class ExportArtifact:
    """Wraps exported file bytes with provenance metadata."""

    def __init__(
        self,
        content: bytes,
        filename: str,
        questionnaire_id: str,
        version: int,
        format: str = "docx",
    ) -> None:
        self.content = content
        self.filename = filename
        self.questionnaire_id = questionnaire_id
        self.version = version
        self.format = format
        self.created_at = datetime.now(tz=timezone.utc)
        self.size_bytes = len(content)

    def save_to(self, directory: Path) -> Path:
        """Save artifact to a directory with provenance. Returns the file path."""
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / self.filename
        path.write_bytes(self.content)
        return path

    def provenance(self) -> dict[str, Any]:
        return {
            "filename": self.filename,
            "questionnaire_id": self.questionnaire_id,
            "version": self.version,
            "format": self.format,
            "size_bytes": self.size_bytes,
            "created_at": self.created_at.isoformat(),
        }


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _add_title_page(doc: Document, qre: Questionnaire) -> None:
    """Add a title page with questionnaire metadata."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research Questionnaire")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()  # spacer

    meta_lines = [
        f"Project: {qre.project_id}",
        f"Methodology: {qre.methodology}",
        f"Version: {qre.version}",
        f"Sections: {len(qre.sections)}",
        f"Questions: {qre.total_questions}",
        f"Estimated LOI: {qre.estimated_loi_minutes} minutes",
    ]
    if qre.context_hash:
        meta_lines.append(f"Context Hash: {qre.context_hash}")

    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def _format_question_type(qt: QuestionType) -> str:
    return {
        QuestionType.SINGLE_SELECT: "Single Select",
        QuestionType.MULTI_SELECT: "Multi Select",
        QuestionType.LIKERT_SCALE: "Likert Scale",
        QuestionType.NUMERIC: "Numeric",
        QuestionType.OPEN_ENDED: "Open-Ended",
        QuestionType.MAXDIFF_TASK: "MaxDiff Task",
        QuestionType.RANKING: "Ranking",
    }.get(qt, qt.value)


def _add_section(doc: Document, section: Section) -> None:
    """Add one section to the document."""
    # Section header
    heading = doc.add_heading(f"Section {section.order + 1}: {section.label}", level=1)

    # Section metadata
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run(f"Type: {section.section_type} | Questions: {len(section.questions)}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Questions
    for q in section.questions:
        _add_question(doc, q)

    doc.add_paragraph()  # spacer between sections


def _add_question(doc: Document, q: Question) -> None:
    """Add one question to the document."""
    # Question ID + variable name header
    id_line = doc.add_paragraph()
    id_run = id_line.add_run(f"[{q.question_id}]  var: {q.var_name}  |  {_format_question_type(q.question_type)}")
    id_run.font.size = Pt(8)
    id_run.font.color.rgb = RGBColor(100, 100, 100)
    id_line.paragraph_format.space_after = Pt(2)

    # Question text
    q_para = doc.add_paragraph(q.question_text)
    q_para.runs[0].bold = True
    q_para.paragraph_format.space_after = Pt(4)

    # Response options
    if q.response_options:
        for opt in q.response_options:
            prefix = "○" if q.question_type == QuestionType.SINGLE_SELECT else "☐"
            term_note = " [TERMINATE]" if opt.terminates else ""
            doc.add_paragraph(
                f"  {prefix}  ({opt.code}) {opt.label}{term_note}",
                style="List Bullet",
            )

    # Scale labels
    if q.scale_points and q.scale_labels:
        scale_text = " | ".join(f"{k}={v}" for k, v in sorted(q.scale_labels.items()))
        scale_para = doc.add_paragraph(f"Scale ({q.scale_points}-point): {scale_text}")
        scale_para.runs[0].font.size = Pt(8)
        scale_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    # Logic notes
    if q.logic:
        logic_para = doc.add_paragraph(f"Logic: {q.logic}")
        logic_para.runs[0].font.size = Pt(8)
        logic_para.runs[0].font.italic = True
        logic_para.runs[0].font.color.rgb = RGBColor(150, 80, 80)


def export_questionnaire_docx(qre: Questionnaire) -> ExportArtifact:
    """Generate a DOCX export from a Questionnaire.

    Returns an ExportArtifact with the file bytes and provenance.
    """
    doc = Document()

    # Title page
    _add_title_page(doc, qre)

    # Sections
    for section in qre.sections:
        _add_section(doc, section)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()

    filename = f"{qre.project_id}_questionnaire_v{qre.version}.docx"

    return ExportArtifact(
        content=content,
        filename=filename,
        questionnaire_id=qre.questionnaire_id,
        version=qre.version,
    )
